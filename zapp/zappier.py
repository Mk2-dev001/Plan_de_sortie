#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script Zapier-like pour connecter Google Drive à WordPress
Surveille le dossier OUTBOX dans Google Drive et crée des articles en brouillon sur WordPress
"""

import os
import json
import time
import logging
from typing import Dict, List, Optional
import requests
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
import io
import docx
from docx.oxml.shared import OxmlElement, qn
import re
from openai import OpenAI
import tempfile
import uuid

# Configuration du logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('zappier.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configuration du client OpenAI
client = OpenAI(api_key="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")

def add_hyperlink(paragraph, text, url):
    """
    Ajoute un hyperlien dans un paragraphe Word.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    return hyperlink

def load_wordpress_db():
    """Charge la base de données WordPress depuis le fichier JSON"""
    try:
        with open('export_wordpress_propre.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Erreur lors du chargement de la base de données WordPress: {str(e)}")
        return []

def analyze_text_with_gpt(text, wp_db, aggressive_mode=True):
    """
    Analyse le texte avec GPT-4 pour détecter les entités importantes.
    
    Args:
        text: Le texte à analyser
        wp_db: La base de données WordPress
        aggressive_mode: Mode agressif pour détecter plus d'entités
    """
    try:
        logger.info(f"Analyse du texte avec GPT-4. Longueur du texte: {len(text)} caractères")
        
        # Vérifier si le texte est trop court pour être analysé
        if len(text.strip()) < 3:  # Seuil plus bas en mode agressif
            logger.info("Texte trop court, pas d'analyse nécessaire")
            return []
        
        examples = []
        for item in wp_db[:5]:
            examples.append(f'{{"text": "{item["title"]}", "url": "{item["link"]}"}}')
        
        examples_str = "\n".join(examples)
        available_titles = [item["title"] for item in wp_db]
        titles_str = "\n".join(available_titles[:50])
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"""Tu es un assistant qui analyse du texte et identifie les entités importantes qui méritent d'être liées.
                IMPORTANT: Tu dois répondre UNIQUEMENT avec un JSON valide, sans aucun autre texte avant ou après.
                
                Voici les titres d'articles disponibles dans la base de données:
                {titles_str}

                Le JSON doit être dans ce format exact:
                {{
                    "entities": [
                        {{
                            "text": "texte original",
                            "url": "url complète de l'article"
                        }}
                    ]
                }}

                Voici des exemples de correspondances:
                {examples_str}

                Règles importantes:
                1. Ne mets aucun texte avant ou après le JSON
                2. Le JSON doit être la seule chose dans ta réponse
                3. Fais correspondre le texte avec les titres de la base de données en utilisant:
                   - Correspondance exacte
                   - Correspondance partielle (si le texte est contenu dans un titre)
                   - Correspondance par mots-clés (si plusieurs mots correspondent)
                4. Détecte les entités suivantes (SOYEZ TRÈS GÉNÉREUX):
                   - Noms propres (personnes, organisations, marques)
                   - Titres d'œuvres (films, livres, séries, jeux vidéo, albums)
                   - Noms de réalisateurs, auteurs, artistes, acteurs, actrices, musiciens
                   - Noms de personnages
                   - Noms de lieux (villes, pays, festivals, salles)
                   - Noms de studios, maisons de production, labels
                   - Noms de technologies, plateformes, services
                   - Noms de genres, mouvements artistiques
                   - Noms de prix, récompenses, événements
                   - Noms de personnalités publiques, influenceurs
                5. Si tu trouves plusieurs correspondances possibles, choisis la plus pertinente
                6. Ne crée pas d'URLs, utilise uniquement celles de la base de données
                7. Si un mot ou une phrase apparaît plusieurs fois, crée un lien pour chaque occurrence
                8. IMPORTANT: Ne crée PAS de liens pour:
                   - Les mots très courts (moins de 2 caractères)
                   - Les mots de liaison très communs (le, la, les, un, une, des, et, ou, mais)
                9. Si aucun lien n'est trouvé, retourne un JSON avec un tableau vide: {{"entities": []}}
                10. SOYEZ TRÈS GÉNÉREUX dans la détection - il vaut mieux avoir trop de liens que pas assez
                11. Incluez même les correspondances partielles et les variations de noms
                12. Détectez les abréviations et les noms complets séparément"""},
                {"role": "user", "content": text}
            ],
            temperature=0.1,
            max_tokens=3000
        )
        
        content = response.choices[0].message.content.strip()
        
        # Log de la réponse brute pour debug
        logger.debug(f"Réponse GPT brute: '{content}'")
        
        # Vérifier si la réponse est vide
        if not content:
            logger.warning("Réponse GPT vide")
            return []
        
        try:
            result = json.loads(content)
            if not isinstance(result, dict) or "entities" not in result:
                logger.warning("Format de réponse invalide")
                return []
            
            valid_entities = []
            for entity in result.get("entities", []):
                # Filtrer les entités non désirées (filtres plus permissifs)
                if any([
                    len(entity["text"]) < 2,  # Seulement les mots très courts
                    entity["text"].lower() in ["le", "la", "les", "un", "une", "des", "et", "ou", "mais"]  # Mots de liaison très basiques
                ]):
                    continue
                
                # Chercher une correspondance dans la base de données (plus généreux)
                best_match = None
                best_score = 0
                
                for wp_item in wp_db:
                    wp_title = wp_item["title"].lower()
                    entity_text = entity["text"].lower()
                    
                    # Correspondance exacte
                    if wp_title == entity_text:
                        best_match = wp_item
                        best_score = 100
                        break
                    # Correspondance partielle (l'entité est dans le titre)
                    elif entity_text in wp_title:
                        score = len(entity_text) / len(wp_title) * 60
                        if score > best_score:
                            best_match = wp_item
                            best_score = score
                    # Correspondance partielle (le titre est dans l'entité)
                    elif wp_title in entity_text:
                        score = len(wp_title) / len(entity_text) * 40
                        if score > best_score:
                            best_match = wp_item
                            best_score = score
                    # Correspondance par mots-clés (au moins 2 mots en commun)
                    else:
                        entity_words = set(entity_text.split())
                        title_words = set(wp_title.split())
                        common_words = entity_words.intersection(title_words)
                        if len(common_words) >= 2:
                            score = len(common_words) / max(len(entity_words), len(title_words)) * 30
                            if score > best_score:
                                best_match = wp_item
                                best_score = score
                
                # Si on a trouvé une correspondance valide (seuil plus bas)
                if best_match and best_score > 5:
                    entity["url"] = best_match["link"]
                    valid_entities.append(entity)
            
            logger.info(f"Entités valides trouvées: {len(valid_entities)}")
            return valid_entities
            
        except json.JSONDecodeError as e:
            logger.error(f"Erreur de parsing JSON: {str(e)}")
            logger.error(f"Contenu reçu: '{content}'")
            return []
            
    except Exception as e:
        logger.error(f"Erreur lors de l'analyse avec GPT: {str(e)}")
        return []

def process_word_document_with_hyperlinks(file_content):
    """
    Traite un document Word pour ajouter des hyperliens intelligemment.
    """
    try:
        logger.info("Début du traitement du document Word avec hyperliens")
        
        wp_db = load_wordpress_db()
        if not wp_db:
            logger.error("Impossible de charger la base de données WordPress")
            return None
            
        # Créer un fichier temporaire pour le document Word
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_input.write(file_content)
        temp_input.close()
        
        # Lire le document Word
        try:
            doc = docx.Document(temp_input.name)
        except Exception as e:
            logger.error(f"Impossible d'ouvrir le document Word: {str(e)}")
            os.unlink(temp_input.name)
            return None
        
        new_doc = docx.Document()
        
        for para in doc.paragraphs:
            new_para = new_doc.add_paragraph()
            doc_text = para.text
            
            if doc_text.strip():
                # NE PAS traiter les paragraphes qui contiennent des métadonnées
                doc_text_stripped = doc_text.strip().upper()
                if (doc_text_stripped.startswith('TITRE :') or 
                    doc_text_stripped.startswith('CATEGORIE :') or 
                    doc_text_stripped.startswith('TAGS :') or 
                    doc_text_stripped.startswith('AUTEUR :') or 
                    doc_text_stripped.startswith('SEO_KEYWORD :') or 
                    doc_text_stripped.startswith('EXCERPT :') or 
                    doc_text_stripped.startswith('CONTENU :')):
                    # C'est une métadonnée, la copier telle quelle SANS AUCUN TRAITEMENT
                    new_para.add_run(doc_text)
                    continue
                
                # Seulement analyser avec GPT si ce n'est PAS une métadonnée
                entities = analyze_text_with_gpt(doc_text, wp_db, aggressive_mode=True)
                current_text = doc_text
                last_pos = 0
                sorted_entities = sorted(entities, key=lambda x: current_text.find(x["text"]))
                
                # Préserver le texte original et ajouter les hyperliens par-dessus
                # Au lieu de remplacer le texte, on va créer des hyperliens qui préservent le texte
                current_pos = 0
                
                for entity in sorted_entities:
                    pos = current_text.find(entity["text"], current_pos)
                    if pos != -1:
                        # Ajouter le texte avant l'entité
                        if pos > current_pos:
                            new_para.add_run(current_text[current_pos:pos])
                        
                        # Créer un hyperlien qui préserve le texte original
                        add_hyperlink(new_para, entity["text"], entity["url"])
                        
                        current_pos = pos + len(entity["text"])
                
                # Ajouter le reste du texte
                if current_pos < len(current_text):
                    new_para.add_run(current_text[current_pos:])
        
        # Sauvegarder le nouveau document
        output_filename = f"processed_{uuid.uuid4()}.docx"
        output_path = os.path.join(tempfile.gettempdir(), output_filename)
        new_doc.save(output_path)
        
        # Nettoyer le fichier temporaire d'entrée
        os.unlink(temp_input.name)
        
        logger.info(f"Document traité avec hyperliens sauvegardé: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Erreur lors du traitement du document avec hyperliens: {str(e)}")
        return None

class GoogleDriveToWordPress:
    def __init__(self, config_file: str = "zappier_config.json"):
        """
        Initialise le connecteur Google Drive vers WordPress
        
        Args:
            config_file: Chemin vers le fichier de configuration
        """
        self.config = self.load_config(config_file)
        self.drive_service = None
        self.wordpress_session = requests.Session()
        self.processed_files = set()
        self.load_processed_files()
        
    def load_config(self, config_file: str) -> Dict:
        """Charge la configuration depuis un fichier JSON"""
        if os.path.exists(config_file):
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        else:
            # Configuration par défaut
            config = {
                "google_drive": {
                    "credentials_file": "credentials.json",
                    "token_file": "token.json",
                    "folder_name": "OUTBOX",
                    "scopes": [
                        "https://www.googleapis.com/auth/drive.readonly",
                        "https://www.googleapis.com/auth/drive.file"
                    ]
                },
                "wordpress": {
                    "site_url": "https://votre-site.com",
                    "username": "votre_username",
                    "password": "votre_password",
                    "api_endpoint": "/wp-json/wp/v2/posts"
                },
                "monitoring": {
                    "check_interval": 300,  # 5 minutes
                    "max_file_size": 10485760  # 10MB
                },
                "content_processing": {
                    "enable_hyperlinks": True,  # Active les hyperliens automatiques
                    "hyperlinks_aggressive": True,  # Mode agressif pour plus d'hyperliens
                    "save_processed_documents": True,  # Sauvegarde les documents traités dans Google Drive
                    "supported_formats": [
                        "application/vnd.google-apps.document",
                        "text/plain",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ]
                }
            }
            self.save_config(config, config_file)
            logger.info(f"Fichier de configuration créé: {config_file}")
            logger.warning("Veuillez configurer vos identifiants dans le fichier de configuration")
            return config
    
    def save_config(self, config: Dict, config_file: str):
        """Sauvegarde la configuration dans un fichier JSON"""
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
    
    def load_processed_files(self):
        """Charge la liste des fichiers déjà traités"""
        if os.path.exists("processed_files.json"):
            with open("processed_files.json", 'r', encoding='utf-8') as f:
                self.processed_files = set(json.load(f))
    
    def save_processed_files(self):
        """Sauvegarde la liste des fichiers traités"""
        with open("processed_files.json", 'w', encoding='utf-8') as f:
            json.dump(list(self.processed_files), f, indent=2)
    
    def authenticate_google_drive(self):
        """Authentification avec Google Drive API"""
        creds = None
        
        # Charge les tokens existants
        if os.path.exists(self.config["google_drive"]["token_file"]):
            creds = Credentials.from_authorized_user_file(
                self.config["google_drive"]["token_file"], 
                self.config["google_drive"]["scopes"]
            )
        
        # Si pas de tokens valides, demande l'authentification
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(
                    self.config["google_drive"]["credentials_file"],
                    self.config["google_drive"]["scopes"]
                )
                creds = flow.run_local_server(port=0)
            
            # Sauvegarde les tokens
            with open(self.config["google_drive"]["token_file"], 'w') as token:
                token.write(creds.to_json())
        
        self.drive_service = build('drive', 'v3', credentials=creds)
        logger.info("Authentification Google Drive réussie")
    
    def set_wordpress_basic_auth(self):
        """Configure l'authentification Basic Auth pour WordPress"""
        import base64
        username = self.config["wordpress"]["username"]
        password = self.config["wordpress"]["password"]
        credentials = f"{username}:{password}"
        token = base64.b64encode(credentials.encode()).decode()
        self.wordpress_session.headers.update({
            "Authorization": f"Basic {token}",
            "Content-Type": "application/json"
        })
    
    def find_folder_id(self, folder_name: str) -> Optional[str]:
        """Trouve l'ID du dossier Google Drive par son nom"""
        try:
            results = self.drive_service.files().list(
                q=f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false",
                spaces='drive',
                fields='files(id, name)'
            ).execute()
            
            files = results.get('files', [])
            if files:
                folder_id = files[0]['id']
                logger.info(f"Dossier trouvé: {folder_name} (ID: {folder_id})")
                return folder_id
            else:
                logger.warning(f"Dossier '{folder_name}' non trouvé")
                return None
                
        except Exception as e:
            logger.error(f"Erreur lors de la recherche du dossier: {e}")
            return None
    
    def get_files_from_folder(self, folder_id: str) -> List[Dict]:
        """Récupère tous les fichiers du dossier spécifié"""
        try:
            results = self.drive_service.files().list(
                q=f"'{folder_id}' in parents and trashed=false",
                spaces='drive',
                fields='files(id, name, mimeType, size, modifiedTime, description)',
                orderBy='modifiedTime desc'
            ).execute()
            
            files = results.get('files', [])
            logger.info(f"{len(files)} fichiers trouvés dans le dossier")
            return files
            
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des fichiers: {e}")
            return []
    
    def download_file_content(self, file_id: str, mime_type: str, file_name: str = "document") -> Optional[str]:
        """Télécharge le contenu d'un fichier Google Drive"""
        try:
            if 'google-apps' in mime_type:
                # Pour les documents Google (Docs, Sheets, etc.)
                if 'document' in mime_type:
                    content = self.drive_service.files().export(
                        fileId=file_id, 
                        mimeType='text/plain'
                    ).execute()
                    return content.decode('utf-8')
                elif 'spreadsheet' in mime_type:
                    content = self.drive_service.files().export(
                        fileId=file_id, 
                        mimeType='text/csv'
                    ).execute()
                    return content.decode('utf-8')
                else:
                    logger.warning(f"Type de fichier Google non supporté: {mime_type}")
                    return None
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Cas d'un fichier Word (.docx)
                request = self.drive_service.files().get_media(fileId=file_id)
                file = io.BytesIO()
                downloader = MediaIoBaseDownload(file, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                file.seek(0)
                
                # Vérifier si les hyperliens sont activés
                if self.config.get("content_processing", {}).get("enable_hyperlinks", False):
                    # Traiter le document avec hyperliens
                    file_content = file.read()
                    processed_file_path = process_word_document_with_hyperlinks(file_content)
                    
                    if processed_file_path:
                        # Convertir le document traité en HTML avec hyperliens
                        html_content = self.convert_docx_to_html_with_hyperlinks(processed_file_path)
                        
                        # Sauvegarder le document traité dans Google Drive si activé
                        if self.config.get("content_processing", {}).get("save_processed_documents", False):
                            try:
                                self.save_processed_document_to_drive(file_id, processed_file_path, file_name)
                            except Exception as e:
                                logger.warning(f"Impossible de sauvegarder le document traité: {e}")
                        
                        # Nettoyer le fichier temporaire traité
                        try:
                            os.unlink(processed_file_path)
                        except:
                            pass
                        
                        logger.info(f"Document Word traité avec hyperliens et converti en HTML: {file_id}")
                        return html_content
                
                # Fallback: lecture normale (avec ou sans hyperliens échoués)
                file.seek(0)
                doc = docx.Document(file)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                return '\n'.join(full_text)
            else:
                # Pour les fichiers normaux
                request = self.drive_service.files().get_media(fileId=file_id)
                file = io.BytesIO()
                downloader = MediaIoBaseDownload(file, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                file.seek(0)
                return file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Erreur lors du téléchargement du fichier {file_id}: {e}")
            return None

    def convert_docx_to_html_with_hyperlinks(self, docx_path: str) -> str:
        """Convertit un document Word avec hyperliens en HTML"""
        try:
            doc = docx.Document(docx_path)
            html_parts = []
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    html_parts.append('<p>&nbsp;</p>')
                    continue
                
                # Traiter chaque run dans le paragraphe pour détecter les hyperliens
                para_html = '<p>'
                
                # Obtenir tous les hyperliens du paragraphe
                hyperlinks = para._element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/officeDocument/2006/wordprocessingml'})
                
                # Créer un mapping des positions des hyperliens
                hyperlink_map = {}
                for hyperlink in hyperlinks:
                    r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if r_id and r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        # Trouver le texte de l'hyperlien
                        for run_elem in hyperlink.findall('.//w:r', {'w': 'http://schemas.openxmlformats.org/officeDocument/2006/wordprocessingml'}):
                            text_elem = run_elem.find('.//w:t', {'w': 'http://schemas.openxmlformats.org/officeDocument/2006/wordprocessingml'})
                            if text_elem is not None:
                                text = text_elem.text or ""
                                if text:
                                    hyperlink_map[text] = url
                
                # Traiter chaque run
                for run in para.runs:
                    run_text = run.text
                    if not run_text:
                        continue
                    
                    # Vérifier si ce run contient un hyperlien
                    has_hyperlink = False
                    for link_text, url in hyperlink_map.items():
                        if link_text in run_text:
                            # Diviser le texte si nécessaire
                            parts = run_text.split(link_text)
                            for i, part in enumerate(parts):
                                if part:
                                    para_html += part
                                if i < len(parts) - 1:  # Pas après le dernier élément
                                    para_html += f'<a href="{url}" target="_blank">{link_text}</a>'
                            has_hyperlink = True
                            break
                    
                    if not has_hyperlink:
                        # Pas d'hyperlien, ajouter le texte
                        para_html += run_text
                
                para_html += '</p>'
                html_parts.append(para_html)
            
            return '\n'.join(html_parts)
            
        except Exception as e:
            logger.error(f"Erreur lors de la conversion en HTML: {str(e)}")
            # Fallback: conversion simple en texte
            doc = docx.Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)


    
    def is_supported_format(self, mime_type: str) -> bool:
        """Vérifie si le format de fichier est supporté"""
        supported_formats = self.config.get("content_processing", {}).get("supported_formats", [
            "application/vnd.google-apps.document",
            "text/plain",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ])
        return mime_type in supported_formats
    
    def parse_content_format(self, content: str) -> Dict:
        """Parse le contenu selon le format spécifique avec métadonnées - FILTRE LES MÉTADONNÉES DU CONTENU FINAL"""
        parsed = {}
        lines = content.split('\n')
        
        # Extraire les métadonnées et construire le contenu filtré
        content_lines = []
        in_content_section = False
        
        for line in lines:
            line_stripped = line.strip()
            
            # Détecte les sections par leurs préfixes
            if line_stripped.upper().startswith('TITRE :'):
                # Extraire le titre - tout ce qui suit "TITRE :"
                title_part = line_stripped[7:].strip()  # Enlever "TITRE :"
                if title_part:
                    parsed['titre'] = title_part
            elif line_stripped.upper().startswith('CATEGORIE :'):
                category_part = line_stripped[11:].strip()  # Enlever "CATEGORIE :"
                if category_part:
                    parsed['categorie'] = category_part
            elif line_stripped.upper().startswith('TAGS :'):
                tags_part = line_stripped[6:].strip()  # Enlever "TAGS :"
                if tags_part:
                    parsed['tags'] = tags_part
            elif line_stripped.upper().startswith('AUTEUR :'):
                author_part = line_stripped[8:].strip()  # Enlever "AUTEUR :"
                if author_part:
                    parsed['auteur'] = author_part
            elif line_stripped.upper().startswith('SEO_KEYWORD :'):
                seo_part = line_stripped[13:].strip()  # Enlever "SEO_KEYWORD :"
                if seo_part:
                    parsed['seo_keyword'] = seo_part
            elif line_stripped.upper().startswith('EXCERPT :'):
                excerpt_part = line_stripped[9:].strip()  # Enlever "EXCERPT :"
                if excerpt_part:
                    parsed['excerpt'] = excerpt_part
            elif line_stripped.upper().startswith('CONTENU :'):
                in_content_section = True
            else:
                # Si on est dans la section CONTENU, ajouter la ligne au contenu
                if in_content_section:
                    content_lines.append(line)
        
        # Construire le contenu final (sans les métadonnées)
        if content_lines:
            parsed['contenu'] = '\n'.join(content_lines).strip()
        else:
            parsed['contenu'] = ""
        
        # Valeurs par défaut pour les métadonnées manquantes
        if not parsed.get('titre'):
            parsed['titre'] = "Sans titre"
        
        return parsed
    
    def extract_file_info(self, file_data: Dict, content: str) -> Dict:
        """Extrait les informations du fichier pour créer l'article"""
        # Vérifier si le contenu est en HTML (contient des balises)
        is_html_content = '<p>' in content or '<a href=' in content
        
        if is_html_content:
            # Pour le contenu HTML, extraire le texte brut pour le parsing des métadonnées
            import re
            # Supprimer les balises HTML pour extraire le texte brut
            text_content = re.sub(r'<[^>]+>', '', content)
            text_content = re.sub(r'&nbsp;', ' ', text_content)
            
            # Parse le contenu selon le format spécifique
            parsed_data = self.parse_content_format(text_content)
            
            # Utiliser le contenu HTML pour le contenu final
            parsed_data['contenu'] = content
        else:
            # Parse le contenu selon le format spécifique
            parsed_data = self.parse_content_format(content)
        
        # Traiter le contenu avec des hyperliens automatiques si activé
        # MAIS seulement sur le contenu principal, pas sur les métadonnées
        if self.config.get("content_processing", {}).get("enable_hyperlinks", False):
            try:
                wp_db = load_wordpress_db()
                if wp_db:
                    # Analyser le contenu pour ajouter des hyperliens
                    # IMPORTANT: Ne traiter que le contenu principal, pas les métadonnées
                    content_with_links = self.add_hyperlinks_to_content(parsed_data['contenu'], wp_db)
                    parsed_data['contenu'] = content_with_links
                    is_html_content = True
                    logger.info("Hyperliens automatiques ajoutés au contenu")
            except Exception as e:
                logger.warning(f"Impossible d'ajouter les hyperliens automatiques: {e}")
        
        # Gère les catégories et tags
        category_name = parsed_data.get("categorie", "Non classé")
        category_id = self.create_or_get_category(category_name)
        
        tag_names = self.parse_tags(parsed_data.get("tags", ""))
        tag_ids = self.create_or_get_tags(tag_names)
        
        file_info = {
            "title": parsed_data.get("titre", file_data.get("name", "Sans titre")),
            "content": parsed_data.get("contenu", "Contenu non disponible"),
            "excerpt": parsed_data.get("excerpt", ""),
            "status": "draft",
            "categories": [category_id],
            "tags": tag_ids,
            "author": parsed_data.get("auteur", ""),
            "seo_keyword": parsed_data.get("seo_keyword", ""),
            "meta": {
                "source_file": file_data.get("name"),
                "file_id": file_data.get("id"),
                "mime_type": file_data.get("mimeType"),
                "modified_time": file_data.get("modifiedTime"),
                "file_size": file_data.get("size", "N/A"),
                "author": parsed_data.get("auteur", ""),
                "seo_keyword": parsed_data.get("seo_keyword", ""),
                "original_category": parsed_data.get("categorie", ""),
                "has_hyperlinks": is_html_content
            }
        }
        
        return file_info
    
    def parse_tags(self, tags_string: str) -> List[str]:
        """Parse la chaîne de tags en liste"""
        if not tags_string:
            return []
        
        # Nettoyer la chaîne de tags
        tags_string = tags_string.strip()
        
        # Gérer les cas où il n'y a que des virgules ou des espaces
        if tags_string in [',', ', ,', ', , ,'] or not tags_string:
            return []
        
        # Sépare les tags par virgules et nettoie
        tags = []
        for tag in tags_string.split(','):
            tag = tag.strip()
            if tag and tag not in ['', ' ']:
                tags.append(tag)
        
        return tags
    
    def add_hyperlinks_to_content(self, content: str, wp_db: List[Dict]) -> str:
        """Ajoute des hyperliens automatiques au contenu texte - PRÉSERVE TOUT LE CONTENU ORIGINAL"""
        try:
            # Si le contenu est déjà en HTML, le traiter par paragraphes
            if '<p>' in content:
                # Diviser le contenu en paragraphes HTML
                import re
                paragraphs = re.split(r'(<p[^>]*>.*?</p>)', content, flags=re.DOTALL)
                result_parts = []
                
                for part in paragraphs:
                    if part.startswith('<p') and part.endswith('</p>'):
                        # Extraire le texte du paragraphe HTML pour analyse
                        text_content = re.sub(r'<[^>]+>', '', part)
                        text_content = re.sub(r'&nbsp;', ' ', text_content)
                        
                        # NE PAS traiter les paragraphes qui contiennent des métadonnées
                        text_stripped = text_content.strip().upper()
                        if (text_stripped.startswith('TITRE :') or 
                            text_stripped.startswith('CATEGORIE :') or 
                            text_stripped.startswith('TAGS :') or 
                            text_stripped.startswith('AUTEUR :') or 
                            text_stripped.startswith('SEO_KEYWORD :') or 
                            text_stripped.startswith('EXCERPT :') or 
                            text_stripped.startswith('CONTENU :')):
                            # C'est une métadonnée, la laisser PARFAITEMENT INTACTE
                            result_parts.append(part)
                            continue
                        
                        # Analyser ce paragraphe
                        entities = analyze_text_with_gpt(text_content, wp_db, aggressive_mode=True)
                        
                        if entities:
                            # Appliquer les hyperliens au paragraphe HTML
                            processed_part = self.apply_hyperlinks_to_html_paragraph(part, entities)
                            result_parts.append(processed_part)
                        else:
                            result_parts.append(part)
                    else:
                        result_parts.append(part)
                
                return ''.join(result_parts)
            else:
                # Contenu texte simple - préserver intégralement
                entities = analyze_text_with_gpt(content, wp_db, aggressive_mode=True)
                
                if not entities:
                    return content
                
                # Trier les entités par longueur (du plus long au plus court pour éviter les conflits)
                entities.sort(key=lambda x: len(x["text"]), reverse=True)
                
                # Remplacer les entités par des hyperliens HTML - PRÉSERVATION TOTALE
                result_content = content
                
                for entity in entities:
                    text = entity["text"]
                    url = entity["url"]
                    
                    # Éviter de créer des liens dans des liens existants
                    if f'<a href="{url}">{text}</a>' not in result_content:
                        # Créer un lien HTML qui préserve le texte original
                        link_html = f'<a href="{url}" target="_blank">{text}</a>'
                        # Utiliser une méthode de remplacement qui préserve le contenu
                        # IMPORTANT: Le texte original est préservé dans le lien HTML
                        result_content = result_content.replace(text, link_html)
                
                return result_content
            
        except Exception as e:
            logger.error(f"Erreur lors de l'ajout d'hyperliens: {e}")
            # En cas d'erreur, retourner le contenu original intact
            return content
    
    def apply_hyperlinks_to_html_paragraph(self, html_paragraph: str, entities: List[Dict]) -> str:
        """Applique les hyperliens à un paragraphe HTML existant"""
        try:
            # Trier les entités par longueur
            entities.sort(key=lambda x: len(x["text"]), reverse=True)
            
            result_html = html_paragraph
            
            for entity in entities:
                text = entity["text"]
                url = entity["url"]
                
                # Éviter de créer des liens dans des liens existants
                if f'<a href="{url}">{text}</a>' not in result_html:
                    # Remplacer le texte par un lien HTML, en préservant les balises existantes
                    import re
                    # Pattern pour remplacer le texte sans affecter les balises HTML
                    pattern = re.compile(re.escape(text), re.IGNORECASE)
                    link_html = f'<a href="{url}" target="_blank">{text}</a>'
                    result_html = pattern.sub(link_html, result_html)
            
            return result_html
            
        except Exception as e:
            logger.error(f"Erreur lors de l'application d'hyperliens au paragraphe HTML: {e}")
            return html_paragraph
    

    
    def create_or_get_category(self, category_name: str) -> int:
        """Crée une catégorie si elle n'existe pas et retourne son ID"""
        try:
            # Vérifie si la catégorie existe déjà
            categories_url = f"{self.config['wordpress']['site_url']}/wp-json/wp/v2/categories"
            response = self.wordpress_session.get(categories_url, params={"search": category_name})
            response.raise_for_status()
            
            categories = response.json()
            for category in categories:
                if category["name"].lower() == category_name.lower():
                    return category["id"]
            
            # Crée la nouvelle catégorie
            category_data = {"name": category_name, "slug": category_name.lower().replace(" ", "-")}
            response = self.wordpress_session.post(categories_url, json=category_data)
            response.raise_for_status()
            
            new_category = response.json()
            logger.info(f"Nouvelle catégorie créée: {category_name} (ID: {new_category['id']})")
            return new_category["id"]
            
        except requests.exceptions.RequestException as e:
            logger.error(f"Erreur lors de la gestion de la catégorie {category_name}: {e}")
            return 1  # Catégorie par défaut
    
    def create_or_get_tags(self, tag_names: List[str]) -> List[int]:
        """Crée les tags manquants et retourne leurs IDs"""
        tag_ids = []
        
        for tag_name in tag_names:
            try:
                # Vérifie si le tag existe déjà
                tags_url = f"{self.config['wordpress']['site_url']}/wp-json/wp/v2/tags"
                response = self.wordpress_session.get(tags_url, params={"search": tag_name})
                response.raise_for_status()
                
                tags = response.json()
                tag_id = None
                
                for tag in tags:
                    if tag["name"].lower() == tag_name.lower():
                        tag_id = tag["id"]
                        break
                
                if tag_id is None:
                    # Crée le nouveau tag
                    tag_data = {"name": tag_name, "slug": tag_name.lower().replace(" ", "-")}
                    response = self.wordpress_session.post(tags_url, json=tag_data)
                    response.raise_for_status()
                    
                    new_tag = response.json()
                    tag_id = new_tag["id"]
                    logger.info(f"Nouveau tag créé: {tag_name} (ID: {tag_id})")
                
                tag_ids.append(tag_id)
                
            except requests.exceptions.RequestException as e:
                logger.error(f"Erreur lors de la gestion du tag {tag_name}: {e}")
        
        return tag_ids
    
    def save_processed_document_to_drive(self, original_file_id: str, processed_file_path: str, original_name: str) -> bool:
        """Sauvegarde le document traité avec hyperliens dans Google Drive"""
        try:
            # Vérifier si les permissions d'écriture sont disponibles
            if "https://www.googleapis.com/auth/drive.file" not in self.config["google_drive"]["scopes"]:
                logger.warning("Permissions d'écriture Google Drive non disponibles, impossible de sauvegarder le document traité")
                return False
            
            # Créer un nom pour le fichier traité
            name_without_ext = os.path.splitext(original_name)[0]
            processed_name = f"{name_without_ext}_avec_hyperliens.docx"
            
            # Trouver le dossier de destination
            folder_id = self.find_folder_id(self.config["google_drive"]["folder_name"])
            if not folder_id:
                logger.error("Impossible de trouver le dossier de destination")
                return False
            
            # Préparer les métadonnées du fichier
            file_metadata = {
                'name': processed_name,
                'parents': [folder_id]
            }
            
            # Vérifier que le fichier traité existe
            if not os.path.exists(processed_file_path):
                logger.error(f"Fichier traité introuvable: {processed_file_path}")
                return False
            
            # Créer le média pour l'upload
            with open(processed_file_path, 'rb') as f:
                file_content = f.read()
            
            media = MediaIoBaseUpload(
                io.BytesIO(file_content),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )
            
            # Uploader le fichier
            file = self.drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()
            
            logger.info(f"Document traité sauvegardé dans Google Drive: {processed_name} (ID: {file.get('id')})")
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la sauvegarde du document traité: {e}")
            # Log plus détaillé pour les erreurs de permissions
            if "insufficientPermissions" in str(e) or "403" in str(e):
                logger.error("Erreur de permissions Google Drive. Vérifiez que les scopes incluent 'drive.file'")
                logger.error("Vous devrez peut-être supprimer le fichier token.json et vous reconnecter")
            return False

    def create_wordpress_post(self, post_data: Dict) -> bool:
        """Crée un article en brouillon sur WordPress"""
        try:
            url = f"{self.config['wordpress']['site_url']}{self.config['wordpress']['api_endpoint']}"
            
            # Prépare les données pour WordPress
            wp_data = {
                "title": post_data["title"],
                "content": post_data["content"],
                "excerpt": post_data["excerpt"],
                "status": post_data["status"],
                "categories": post_data["categories"],
                "tags": post_data["tags"],
                "content_format": "html"  # Indique que le contenu est en HTML
            }
            
            # Ajoute l'auteur si disponible
            if post_data.get("author"):
                wp_data["author_name"] = post_data["author"]
            
            response = self.wordpress_session.post(url, json=wp_data)
            response.raise_for_status()
            
            post_id = response.json().get("id")
            logger.info(f"Article créé avec succès (ID: {post_id}): {post_data['title']}")
            
            # Ajoute les métadonnées personnalisées
            if post_data.get("meta"):
                meta_url = f"{url}/{post_id}/meta"
                for key, value in post_data["meta"].items():
                    meta_data = {"key": f"drive_{key}", "value": str(value)}
                    try:
                        self.wordpress_session.post(meta_url, json=meta_data)
                    except requests.exceptions.RequestException as meta_error:
                        logger.warning(f"Impossible d'ajouter la métadonnée {key}: {meta_error}")
            
            # Ajoute les métadonnées SEO si disponibles
            if post_data.get("seo_keyword"):
                seo_meta_url = f"{url}/{post_id}/meta"
                seo_data = {"key": "seo_focus_keyword", "value": post_data["seo_keyword"]}
                try:
                    self.wordpress_session.post(seo_meta_url, json=seo_data)
                except requests.exceptions.RequestException as seo_error:
                    logger.warning(f"Impossible d'ajouter le mot-clé SEO: {seo_error}")
            
            return True
            
        except requests.exceptions.RequestException as e:
            logger.error(f"Erreur lors de la création de l'article: {e}")
            return False
    
    def process_new_files(self):
        """Traite les nouveaux fichiers du dossier Google Drive"""
        folder_name = self.config["google_drive"]["folder_name"]
        folder_id = self.find_folder_id(folder_name)
        
        if not folder_id:
            logger.error(f"Impossible de trouver le dossier '{folder_name}'")
            return
        
        files = self.get_files_from_folder(folder_id)
        
        for file_data in files:
            file_id = file_data["id"]
            
            # Vérifie si le fichier a déjà été traité
            if file_id in self.processed_files:
                continue
            
            # Vérifie la taille du fichier
            file_size = int(file_data.get("size", 0))
            if file_size > self.config["monitoring"]["max_file_size"]:
                logger.warning(f"Fichier trop volumineux ignoré: {file_data['name']}")
                continue
            
            # Vérifie le format du fichier
            if not self.is_supported_format(file_data["mimeType"]):
                logger.warning(f"Format de fichier non supporté ignoré: {file_data['name']} ({file_data['mimeType']})")
                continue
            
            logger.info(f"Traitement du fichier: {file_data['name']}")
            
            # Télécharge le contenu du fichier
            content = self.download_file_content(file_id, file_data["mimeType"], file_data["name"])
            
            if content:
                # Extrait les informations pour l'article
                post_data = self.extract_file_info(file_data, content)
                
                # Crée l'article sur WordPress
                if self.create_wordpress_post(post_data):
                    self.processed_files.add(file_id)
                    self.save_processed_files()
                    logger.info(f"Fichier traité avec succès: {file_data['name']}")
                else:
                    logger.error(f"Échec du traitement: {file_data['name']}")
            else:
                logger.warning(f"Impossible de récupérer le contenu: {file_data['name']}")
    
    def run(self):
        """Lance le processus de surveillance continue"""
        logger.info("Démarrage du script Zapier-like")
        
        # Authentification
        try:
            self.authenticate_google_drive()
            self.set_wordpress_basic_auth()
        except Exception as e:
            logger.error(f"Erreur d'authentification: {e}")
            return
        
        # Boucle de surveillance
        check_interval = self.config["monitoring"]["check_interval"]
        
        while True:
            try:
                logger.info("Vérification des nouveaux fichiers...")
                self.process_new_files()
                
                logger.info(f"Prochaine vérification dans {check_interval} secondes")
                time.sleep(check_interval)
                
            except KeyboardInterrupt:
                logger.info("Arrêt du script demandé par l'utilisateur")
                break
            except Exception as e:
                logger.error(f"Erreur dans la boucle principale: {e}")
                time.sleep(60)  # Attendre 1 minute en cas d'erreur

def main():
    """Fonction principale"""
    print("=== Script Zapier-like Google Drive vers WordPress ===")
    print("Ce script surveille votre dossier Google Drive 'OUTBOX'")
    print("et crée automatiquement des articles en brouillon sur WordPress")
    print()
    
    # Vérifie si le fichier de configuration existe
    config_file = "zappier_config.json"
    if not os.path.exists(config_file):
        print("⚠️  Fichier de configuration non trouvé.")
        print("Un fichier de configuration par défaut sera créé.")
        print("Veuillez le modifier avec vos identifiants avant de relancer le script.")
        print()
    
    # Lance le script
    connector = GoogleDriveToWordPress(config_file)
    connector.run()

if __name__ == "__main__":
    main()
