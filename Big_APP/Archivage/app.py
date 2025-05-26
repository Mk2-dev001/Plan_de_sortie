import streamlit as st
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
import re
from openai import OpenAI
import json
import logging
from datetime import datetime

# Configuration du logging
logging.basicConfig(
    filename=f'processing_logs_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log',
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Configuration du client OpenAI
client = OpenAI(api_key="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")

def add_hyperlink(paragraph, text, url):
    """
    Ajoute un hyperlien dans un paragraphe Word.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object
    new_run = OxmlElement('w:r')

    # Create a new run properties object
    rPr = OxmlElement('w:rPr')

    # Add color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Join all the xml elements together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    return hyperlink

# Charger la base de données WordPress
def load_wordpress_db():
    try:
        with open('export_wordpress_propre.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        st.error(f"Erreur lors du chargement de la base de données WordPress:{str(e)}")
        return []

def analyze_text_with_gpt(text, wp_db):
    """
    Analyse le texte avec GPT-4 pour détecter les entités importantes.
    """
    try:
        logging.info(f"Analyse du texte avec GPT-4. Longueur du texte: {len(text)} caractères")
        
        # Préparer les exemples de la base de données pour le prompt
        examples = []
        for item in wp_db[:5]:  # Augmenté à 5 exemples
            examples.append(f'{{"text": "{item["title"]}", "url": "{item["link"]}"}}')
        
        examples_str = "\n".join(examples)
        logging.debug(f"Exemples préparés: {examples_str}")
        
        # Créer une liste des titres disponibles pour aider GPT
        available_titles = [item["title"] for item in wp_db]
        titles_str = "\n".join(available_titles[:50])  # Augmenté à 50 titres
        logging.debug(f"Titres disponibles: {titles_str}")
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"""Tu es un assistant qui analyse du texte et identifie les entités importantes qui méritent d'être liées.
                IMPORTANT: Tu dois répondre UNIQUEMENT avec un JSON valide, sans aucun autre texte avant ou après.
                
                Voici les titres d'articles disponibles dans la base de données:
                {titles_str}

                Le JSON doit être dans ce format exact:
                {{
                    "entities": [
                        {{
                            "text": "texte original",
                            "url": "url complète de l'article"
                        }}
                    ]
                }}

                Voici des exemples de correspondances:
                {examples_str}

                Règles importantes:
                1. Ne mets aucun texte avant ou après le JSON
                2. Le JSON doit être la seule chose dans ta réponse
                3. Fais correspondre le texte avec les titres de la base de données en utilisant:
                   - Correspondance exacte
                   - Correspondance partielle (si le texte est contenu dans un titre)
                4. Détecte les entités suivantes:
                   - Noms propres (personnes, organisations)
                   - Titres d'œuvres (films, livres, etc.)
                   - Noms de réalisateurs, auteurs, artistes, acteurs, actrices
                5. Si tu trouves plusieurs correspondances possibles, choisis la plus pertinente ( privligie la longeur de l'article )
                6. Ne crée pas d'URLs, utilise uniquement celles de la base de données
                7. Si un mot ou une phrase apparaît plusieurs fois, crée un lien pour chaque occurrence
                8. IMPORTANT: Ne crée PAS de liens pour:
                   - Les titres d'articles
                   - Les dates (années, jours, mois)*
                   - Les lieux (lieux, villes, pays)
                   - Les numéros de version ou de partie"""},
                {"role": "user", "content": text}
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        content = response.choices[0].message.content
        logging.debug(f"Réponse GPT reçue: {content}")
        
        content = content.strip()
        
        try:
            result = json.loads(content)
            if not isinstance(result, dict) or "entities" not in result:
                logging.warning("Format de réponse invalide")
                return []
            
            valid_entities = []
            for entity in result.get("entities", []):
                logging.debug(f"Traitement de l'entité: {entity}")
                
                # Vérifier si l'entité n'est pas un titre ou une date
                if not any([
                    entity["text"].startswith("CANNES"),  # Exclure les titres commençant par CANNES
                    re.match(r'\d{4}', entity["text"]),  # Exclure les années
                    re.match(r'Partie \d+', entity["text"]),  # Exclure les numéros de partie
                    entity["text"].startswith("Mission : Impossible")  # Exclure les titres de films
                ]):
                    # Recherche plus flexible des correspondances
                    for wp_item in wp_db:
                        # Correspondance exacte
                        if wp_item["title"].lower() == entity["text"].lower():
                            entity["url"] = wp_item["link"]
                            valid_entities.append(entity)
                            logging.info(f"Entité valide trouvée (exacte): {entity}")
                            break
                        # Correspondance partielle
                        elif entity["text"].lower() in wp_item["title"].lower():
                            entity["url"] = wp_item["link"]
                            valid_entities.append(entity)
                            logging.info(f"Entité valide trouvée (partielle): {entity}")
                            break
            
            logging.info(f"Nombre total d'entités valides trouvées: {len(valid_entities)}")
            return valid_entities
            
        except json.JSONDecodeError as e:
            logging.error(f"Erreur de parsing JSON: {str(e)}")
            logging.error(f"Contenu reçu: {content}")
            return []
            
    except Exception as e:
        logging.error(f"Erreur lors de l'analyse avec GPT: {str(e)}")
        return []

def process_word_document(file):
    """
    Traite un document Word pour ajouter des hyperliens intelligemment.
    """
    try:
        logging.info("Début du traitement du document Word")
        
        # Charger la base de données WordPress
        wp_db = load_wordpress_db()
        if not wp_db:
            logging.error("Impossible de charger la base de données WordPress")
            return None
            
        logging.info(f"Base de données WordPress chargée avec {len(wp_db)} articles")
        
        # Lire le document Word
        doc = docx.Document(file)
        logging.info(f"Document Word chargé avec {len(doc.paragraphs)} paragraphes")
        
        # Créer un nouveau document
        new_doc = docx.Document()
        
        # Traiter chaque paragraphe
        for i, para in enumerate(doc.paragraphs):
            logging.debug(f"Traitement du paragraphe {i+1}")
            new_para = new_doc.add_paragraph()
            
            try:
                new_para.alignment = para.alignment
            except:
                logging.warning(f"Impossible de copier l'alignement du paragraphe {i+1}")
            
            doc_text = para.text
            if doc_text.strip():
                logging.debug(f"Analyse du texte du paragraphe {i+1}: {doc_text[:100]}...")
                entities = analyze_text_with_gpt(doc_text, wp_db)
                logging.info(f"Entités trouvées dans le paragraphe {i+1}: {len(entities)}")
                
                # Préserver le texte original et ajouter les hyperliens
                current_text = doc_text
                last_pos = 0
                
                # Trier les entités par position dans le texte pour éviter les chevauchements
                sorted_entities = sorted(entities, key=lambda x: current_text.find(x["text"]))
                
                for entity in sorted_entities:
                    pos = current_text.find(entity["text"], last_pos)
                    if pos != -1:
                        # Ajouter le texte avant l'entité
                        if pos > last_pos:
                            new_para.add_run(current_text[last_pos:pos])
                        
                        # Ajouter l'hyperlien
                        add_hyperlink(new_para, entity["text"], entity["url"])
                        last_pos = pos + len(entity["text"])
                
                # Ajouter le reste du texte
                if last_pos < len(current_text):
                    new_para.add_run(current_text[last_pos:])
                
                # Copier le formatage original
                for run in para.runs:
                    try:
                        new_run = new_para.runs[-1]
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                    except:
                        logging.warning(f"Impossible de copier le formatage du run dans le paragraphe {i+1}")
        
        # Sauvegarder le nouveau document
        docx_buffer = io.BytesIO()
        new_doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        logging.info("Document traité avec succès")
        return docx_buffer
        
    except Exception as e:
        logging.error(f"Erreur lors du traitement du document: {str(e)}")
        return None

def main():
    """
    Interface principale de l'application Streamlit.
    """
    logging.info("Démarrage de l'application")
    st.title("🔗 Traitement de Documents Word")
    st.write("Déposez votre document Word pour ajouter des hyperliens automatiquement")
    
    # Informations sur l'application
    with st.expander("ℹ️ Informations"):
        st.write("""
        Cette application:
        - Utilise l'IA (GPT-4) pour détecter intelligemment les entités importantes dans votre texte 
        - Crée des hyperliens vers les articles de votre base WordPress
        - Préserve le formatage original du document
        - Génère un nouveau document avec les hyperliens ajoutés
        """)
    
    # Upload du fichier
    uploaded_file = st.file_uploader(
        "Choisissez un fichier Word", 
        type=['docx'],
        help="Seuls les fichiers .docx sont supportés"
    )
    
    if uploaded_file is not None:
        # Afficher les informations du fichier
        st.info(f"📄 Fichier: {uploaded_file.name} ({uploaded_file.size} bytes)")
        
        # Bouton pour traiter le document
        if st.button("🔄 Traiter le document", type="primary"):
            with st.spinner("Traitement en cours..."):
                try:
                    # Traiter le document
                    processed_doc = process_word_document(uploaded_file)
                    
                    if processed_doc is not None:
                        # Afficher le bouton de téléchargement
                        st.download_button(
                            label="📥 Télécharger le document modifié",
                            data=processed_doc,
                            file_name=f"modifie_{uploaded_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("✅ Document traité avec succès!")
                        st.balloons()
                    
                except Exception as e:
                    st.error(f"❌ Une erreur s'est produite: {str(e)}")
                    st.write("Veuillez vérifier que votre fichier est un document Word valide (.docx)")

if __name__ == "__main__":
    main()