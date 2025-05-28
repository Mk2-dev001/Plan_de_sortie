import os
import tempfile
import requests
from flask import Flask, request, jsonify, send_file
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
import re
from openai import OpenAI
import json
import logging
from datetime import datetime
import uuid
import mimetypes

# Configuration du logging
logging.basicConfig(
    filename=f'api_logs_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log',
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Configuration du client OpenAI
client = OpenAI(api_key="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")

# Création de l'application Flask
app = Flask(__name__)

# Dossier temporaire pour stocker les fichiers
TEMP_DIR = tempfile.mkdtemp()
# Dossier pour les fichiers de debug
DEBUG_DIR = os.path.join(TEMP_DIR, 'debug')
os.makedirs(DEBUG_DIR, exist_ok=True)
# Dictionnaire pour stocker les fichiers traités
processed_files = {}

def add_hyperlink(paragraph, text, url):
    """
    Ajoute un hyperlien dans un paragraphe Word.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    return hyperlink

def load_wordpress_db():
    try:
        with open('export_wordpress_propre.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logging.error(f"Erreur lors du chargement de la base de données WordPress:{str(e)}")
        return []

def analyze_text_with_gpt(text, wp_db):
    """
    Analyse le texte avec GPT-4 pour détecter les entités importantes.
    """
    try:
        logging.info(f"Analyse du texte avec GPT-4. Longueur du texte: {len(text)} caractères")
        
        examples = []
        for item in wp_db[:5]:
            examples.append(f'{{"text": "{item["title"]}", "url": "{item["link"]}"}}')
        
        examples_str = "\n".join(examples)
        available_titles = [item["title"] for item in wp_db]
        titles_str = "\n".join(available_titles[:50])
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"""Tu es un assistant qui analyse du texte et identifie les entités importantes qui méritent d'être liées.
                IMPORTANT: Tu dois répondre UNIQUEMENT avec un JSON valide, sans aucun autre texte avant ou après.
                
                Voici les titres d'articles disponibles dans la base de données:
                {titles_str}

                Le JSON doit être dans ce format exact:
                {{
                    "entities": [
                        {{
                            "text": "texte original",
                            "url": "url complète de l'article"
                        }}
                    ]
                }}

                Voici des exemples de correspondances:
                {examples_str}

                Règles importantes:
                1. Ne mets aucun texte avant ou après le JSON
                2. Le JSON doit être la seule chose dans ta réponse
                3. Fais correspondre le texte avec les titres de la base de données en utilisant:
                   - Correspondance exacte
                   - Correspondance partielle (si le texte est contenu dans un titre)
                4. Détecte les entités suivantes:
                   - Noms propres (personnes, organisations)
                   - Titres d'œuvres (films, livres, etc.)
                   - Noms de réalisateurs, auteurs, artistes, acteurs, actrices
                5. Si tu trouves plusieurs correspondances possibles, choisis la plus pertinente ( privligie la longeur de l'article )
                6. Ne crée pas d'URLs, utilise uniquement celles de la base de données
                7. Si un mot ou une phrase apparaît plusieurs fois, crée un lien pour chaque occurrence
                8. IMPORTANT: Ne crée PAS de liens pour:
                   - Les titres d'articles
                   - Les dates (années, jours, mois)*
                   - Les lieux (lieux, villes, pays)
                   - Les numéros de version ou de partie"""},
                {"role": "user", "content": text}
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        content = response.choices[0].message.content.strip()
        
        try:
            result = json.loads(content)
            if not isinstance(result, dict) or "entities" not in result:
                logging.warning("Format de réponse invalide")
                return []
            
            valid_entities = []
            for entity in result.get("entities", []):
                if not any([
                    entity["text"].startswith("CANNES"),
                    re.match(r'\d{4}', entity["text"]),
                    re.match(r'Partie \d+', entity["text"]),
                    entity["text"].startswith("Mission : Impossible")
                ]):
                    for wp_item in wp_db:
                        if wp_item["title"].lower() == entity["text"].lower():
                            entity["url"] = wp_item["link"]
                            valid_entities.append(entity)
                            break
                        elif entity["text"].lower() in wp_item["title"].lower():
                            entity["url"] = wp_item["link"]
                            valid_entities.append(entity)
                            break
            
            return valid_entities
            
        except json.JSONDecodeError as e:
            logging.error(f"Erreur de parsing JSON: {str(e)}")
            return []
            
    except Exception as e:
        logging.error(f"Erreur lors de l'analyse avec GPT: {str(e)}")
        return []

def download_file_from_url(url):
    """
    Télécharge un fichier depuis une URL avec des headers appropriés et vérifie le type de contenu.
    """
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    try:
        response = requests.get(url, headers=headers, stream=True)
        response.raise_for_status()
        
        # Log des headers de réponse
        logging.info(f"Response headers: {dict(response.headers)}")
        
        # Vérification du Content-Type
        content_type = response.headers.get('Content-Type', '').lower()
        if 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' not in content_type:
            logging.warning(f"Unexpected Content-Type: {content_type}")
        
        # Sauvegarde du fichier brut pour debug
        debug_filename = f"debug_received_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        debug_path = os.path.join(DEBUG_DIR, debug_filename)
        
        with open(debug_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        
        logging.info(f"Debug file saved to: {debug_path}")
        
        # Vérification de la taille du fichier
        file_size = os.path.getsize(debug_path)
        if file_size == 0:
            raise ValueError("Downloaded file is empty")
        
        logging.info(f"Downloaded file size: {file_size} bytes")
        
        return debug_path
        
    except requests.exceptions.RequestException as e:
        logging.error(f"Download error: {str(e)}")
        raise

def verify_docx_file(file_path):
    """
    Vérifie si le fichier est un document Word valide.
    """
    try:
        doc = docx.Document(file_path)
        # Vérification basique du contenu
        if len(doc.paragraphs) == 0:
            logging.warning("Document has no paragraphs")
        return True
    except Exception as e:
        logging.error(f"Document verification failed: {str(e)}")
        return False

def process_word_document(file_content):
    """
    Traite un document Word pour ajouter des hyperliens intelligemment.
    """
    try:
        logging.info("Début du traitement du document Word")
        
        wp_db = load_wordpress_db()
        if not wp_db:
            logging.error("Impossible de charger la base de données WordPress")
            return None
            
        # Créer un fichier temporaire pour le document Word
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_input.write(file_content)
        temp_input.close()
        
        # Vérification du document avant traitement
        if not verify_docx_file(temp_input.name):
            raise ValueError("Invalid or corrupted Word document")
        
        # Lire le document Word avec gestion d'erreur explicite
        try:
            doc = docx.Document(temp_input.name)
        except Exception as e:
            logging.error(f"Failed to open Word document: {str(e)}")
            raise ValueError(f"Failed to open Word document: {str(e)}")
        
        new_doc = docx.Document()
        
        for i, para in enumerate(doc.paragraphs):
            new_para = new_doc.add_paragraph()
            try:
                new_para.alignment = para.alignment
            except:
                logging.warning(f"Impossible de copier l'alignement du paragraphe {i+1}")
            
            doc_text = para.text
            if doc_text.strip():
                entities = analyze_text_with_gpt(doc_text, wp_db)
                current_text = doc_text
                last_pos = 0
                sorted_entities = sorted(entities, key=lambda x: current_text.find(x["text"]))
                
                for entity in sorted_entities:
                    pos = current_text.find(entity["text"], last_pos)
                    if pos != -1:
                        if pos > last_pos:
                            new_para.add_run(current_text[last_pos:pos])
                        add_hyperlink(new_para, entity["text"], entity["url"])
                        last_pos = pos + len(entity["text"])
                
                if last_pos < len(current_text):
                    new_para.add_run(current_text[last_pos:])
                
                for run in para.runs:
                    try:
                        new_run = new_para.runs[-1]
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                    except:
                        logging.warning(f"Impossible de copier le formatage du run dans le paragraphe {i+1}")
        
        # Sauvegarder le nouveau document
        output_filename = f"processed_{uuid.uuid4()}.docx"
        output_path = os.path.join(TEMP_DIR, output_filename)
        new_doc.save(output_path)
        
        # Nettoyer le fichier temporaire d'entrée
        os.unlink(temp_input.name)
        
        return output_filename
        
    except Exception as e:
        logging.error(f"Erreur lors du traitement du document: {str(e)}")
        return None

@app.route('/process', methods=['POST'])
def process_document():
    try:
        data = request.get_json()
        if not data or 'file_url' not in data or 'file_name' not in data:
            return jsonify({"error": "Missing required fields"}), 400
        
        file_url = data['file_url']
        file_name = data['file_name']
        
        # Télécharger le fichier depuis Google Drive avec debugging
        try:
            debug_file_path = download_file_from_url(file_url)
            with open(debug_file_path, 'rb') as f:
                file_content = f.read()
        except Exception as e:
            error_msg = f"Failed to download file: {str(e)}"
            logging.error(error_msg)
            return jsonify({"error": error_msg}), 400
        
        # Traiter le document
        try:
            output_filename = process_word_document(file_content)
            if not output_filename:
                return jsonify({"error": "Document processing failed"}), 500
        except ValueError as e:
            error_msg = f"Document processing error: {str(e)}"
            logging.error(error_msg)
            return jsonify({"error": error_msg}), 500
        except Exception as e:
            error_msg = f"Unexpected error during processing: {str(e)}"
            logging.error(error_msg)
            return jsonify({"error": error_msg}), 500
        
        # Stocker le nom du fichier traité
        processed_files[output_filename] = file_name
        
        # Générer l'URL de téléchargement
        download_url = f"/download/{output_filename}"
        
        return jsonify({
            "status": "done",
            "download_url": download_url
        })
        
    except Exception as e:
        error_msg = f"Request processing error: {str(e)}"
        logging.error(error_msg)
        return jsonify({"error": error_msg}), 500

@app.route('/download/<filename>')
def download_file(filename):
    if filename not in processed_files:
        return jsonify({"error": "File not found"}), 404
    
    file_path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=processed_files[filename]
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True) 