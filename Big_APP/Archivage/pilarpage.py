import streamlit as st
import json
from openai import OpenAI
import logging
from datetime import datetime
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app import process_word_document, load_wordpress_db
import io
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import Counter
from textblob import TextBlob
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd

# Téléchargement des ressources NLTK nécessaires
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Chargement du modèle spaCy pour l'analyse sémantique
try:
    nlp = spacy.load('fr_core_news_md')
except OSError:
    # Si le modèle n'est pas installé, on utilise le modèle français de base
    nlp = spacy.load('fr_core_news_sm')

def calculate_readability_score(text):
    """Calcule le score de lisibilité Flesch-Kincaid adapté au français."""
    sentences = sent_tokenize(text)
    words = word_tokenize(text)
    syllables = sum(len(re.findall(r'[aeiouy]+', word.lower())) for word in words)
    
    if not sentences:
        return 0
        
    avg_sentence_length = len(words) / len(sentences)
    avg_syllables_per_word = syllables / len(words)
    
    # Formule adaptée pour le français
    score = 206.835 - (1.015 * avg_sentence_length) - (84.6 * avg_syllables_per_word)
    return max(0, min(100, score))

def analyze_seo_content(text, main_keyword):
    """Analyse la qualité SEO du contenu généré."""
    # Tokenization et nettoyage
    words = word_tokenize(text.lower())
    stop_words = set(stopwords.words('french'))
    words = [word for word in words if word.isalnum() and word not in stop_words]
    
    # Calcul de la densité du mot-clé
    main_keyword_count = text.lower().count(main_keyword.lower())
    keyword_density = (main_keyword_count / len(words)) * 100 if words else 0
    
    # Analyse des phrases longues
    sentences = sent_tokenize(text)
    long_sentences = [s for s in sentences if len(word_tokenize(s)) > 25]
    
    # Détection des expressions vagues
    vague_expressions = [
        "de plus en plus", "le monde change", "il est important de noter que",
        "il faut savoir que", "il est essentiel de", "il est crucial de",
        "il est primordial de", "il est nécessaire de", "il est fondamental de"
    ]
    found_vague = [expr for expr in vague_expressions if expr in text.lower()]
    
    # Calcul du score de lisibilité
    readability_score = calculate_readability_score(text)
    
    # Détection des répétitions lexicales
    word_freq = Counter(words)
    repetitions = {word: count for word, count in word_freq.items() if count > 5}
    
    return {
        "keyword_density": round(keyword_density, 2),
        "readability_score": round(readability_score, 2),
        "long_sentences": len(long_sentences),
        "vague_expressions": found_vague,
        "repetitions": repetitions,
        "total_words": len(words)
    }

def suggest_internal_links(section_content, wp_db, main_keyword):
    """Suggère des liens internes pertinents pour une section donnée."""
    # Préparation du contenu pour l'analyse
    doc = nlp(section_content)
    section_keywords = [token.text.lower() for token in doc if not token.is_stop and token.is_alpha]
    
    # Création d'une matrice TF-IDF pour la comparaison
    vectorizer = TfidfVectorizer(stop_words='french')
    wp_titles = [item["title"] for item in wp_db]
    wp_titles.append(section_content)  # Ajouter le contenu de la section pour comparaison
    
    try:
        tfidf_matrix = vectorizer.fit_transform(wp_titles)
        # Calculer la similarité avec le contenu de la section
        section_vector = tfidf_matrix[-1]
        similarities = cosine_similarity(section_vector, tfidf_matrix[:-1])[0]
        
        # Obtenir les 2 articles les plus pertinents
        top_indices = similarities.argsort()[-2:][::-1]
        suggested_links = []
        
        for idx in top_indices:
            if similarities[idx] > 0.1:  # Seuil de similarité minimum
                suggested_links.append({
                    "title": wp_titles[idx],
                    "relevance_score": round(similarities[idx] * 100, 2)
                })
        
        return suggested_links
    except Exception as e:
        logging.error(f"Erreur lors de la suggestion des liens internes: {str(e)}")
        return []

def analyze_semantic_fields(text, wp_db):
    """Analyse les champs lexicaux dominants du texte."""
    # Tokenization et nettoyage
    doc = nlp(text)
    words = [token.text.lower() for token in doc if not token.is_stop and token.is_alpha]
    
    # Calcul des fréquences
    word_freq = Counter(words)
    top_keywords = word_freq.most_common(10)
    
    # Regroupement sémantique
    semantic_groups = {}
    for word, freq in top_keywords:
        word_doc = nlp(word)
        for other_word, other_freq in top_keywords:
            if word != other_word:
                other_doc = nlp(other_word)
                similarity = word_doc.similarity(other_doc)
                if similarity > 0.5:  # Seuil de similarité
                    if word not in semantic_groups:
                        semantic_groups[word] = []
                    semantic_groups[word].append((other_word, similarity))
    
    # Comparaison avec les titres WordPress
    wp_titles = " ".join([item["title"] for item in wp_db])
    wp_doc = nlp(wp_titles)
    
    semantic_relevance = {}
    for word, freq in top_keywords:
        word_doc = nlp(word)
        relevance = word_doc.similarity(wp_doc)
        semantic_relevance[word] = round(relevance * 100, 2)
    
    return {
        "top_keywords": top_keywords,
        "semantic_groups": semantic_groups,
        "semantic_relevance": semantic_relevance
    }

# Configuration du logging
log_filename = f'pillar_page_logs_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
logging.basicConfig(
    filename=log_filename,
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Configuration du client OpenAI
client = OpenAI(api_key="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")

def get_latest_log_content():
    """Récupère le contenu du dernier fichier de log."""
    try:
        if os.path.exists(log_filename):
            with open(log_filename, 'r', encoding='utf-8') as f:
                return f.read()
        return "Aucun log disponible pour le moment."
    except Exception as e:
        return f"Erreur lors de la lecture des logs: {str(e)}"

def generate_pillar_page(topic, wp_db):
    """Génère une page pilier optimisée SEO sur un sujet donné."""
    try:
        logging.info(f"Génération d'une page pilier sur le sujet: {topic}")
        
        if not wp_db:
            logging.error("La base de données WordPress est vide")
            return None
            
        # Extraire les titres disponibles pour l'inspiration
        available_titles = [item["title"] for item in wp_db]
        titles_str = "\n".join(available_titles[:50])  # Limiter à 50 titres pour le contexte
        
        logging.info("Envoi de la requête à GPT-4...")
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"""Tu es un expert SEO et créateur de contenu qui génère des pages pilier optimisées.
                Ta tâche est de créer une page pilier complète et optimisée SEO sur le sujet demandé.
                
                Articles disponibles pour inspiration:
                {titles_str}
                
                Règles de création:
                1. Structure de texte brut pour Word:
                   - Titre 1 pour le titre principal
                   - Titre 2 pour les sections principales
                   - Titre 3 pour les sous-sections
                
                2. Optimisation SEO:
                   - Titre principal accrocheur et optimisé avec le mot-clé exact
                   - Introduction captivante (150-200 mots) avec mots-clés naturels
                   - Structure hiérarchique claire
                   - Paragraphes courts et aérés
                   - Utilisation naturelle des mots-clés
                
                3. Contenu:
                   - Long et détaillé (minimum 1500 mots)
                   - Professionnel et engageant
                   - Basé uniquement sur les sujets des articles disponibles
                   - Sans balises de lien ou formatage
                   - Avec des listes à puces quand pertinent
                   - Section FAQ obligatoire avec 4-5 questions fréquentes
                
                4. Style et ton:
                   - Informatif et accessible
                   - Formulations engageantes
                   - Pas de contenu superflu
                   - Priorité à la clarté et la lisibilité
                   - Rythme dynamique
                
                5. Format de réponse:
                {{
                    "title": "Titre principal optimisé SEO",
                    "content": "Introduction de 150-200 mots",
                    "sections": [
                        {{
                            "title": "Titre de la section",
                            "level": 2,
                            "content": "Contenu de la section"
                        }},
                        {{
                            "title": "Titre de la sous-section",
                            "level": 3,
                            "content": "Contenu de la sous-section"
                        }},
                        {{
                            "title": "FAQ",
                            "level": 2,
                            "content": "4-5 questions fréquentes avec réponses détaillées"
                        }}
                    ]
                }}
                
                Important: 
                - Ne pas inclure de balises de lien ou de formatage dans le contenu
                - Les liens seront ajoutés automatiquement par un programme externe
                - Assurez-vous que chaque section est séparée par une ligne vide
                - Utilisez des mots-clés pertinents qui correspondent aux titres d'articles disponibles"""},
                {"role": "user", "content": f"Crée une page pilier optimisée SEO sur le sujet: {topic}"}
            ],
            temperature=0.7,
            max_tokens=2500
        )
        
        content = response.choices[0].message.content
        logging.debug(f"Réponse GPT reçue: {content}")
        
        try:
            # Nettoyer le contenu JSON avant de le parser
            cleaned_content = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', content)
            cleaned_content = re.sub(r'^\s+', '', cleaned_content, flags=re.MULTILINE)
            
            result = json.loads(cleaned_content)
            logging.info("Parsing JSON réussi")
            
            # Vérification de la structure du résultat
            required_fields = ["title", "content", "sections"]
            missing_fields = [field for field in required_fields if field not in result]
            
            if missing_fields:
                logging.error(f"Champs manquants dans la réponse: {missing_fields}")
                return None
            
            # Analyse SEO du contenu
            full_content = result['content'] + "\n\n" + "\n\n".join(section['content'] for section in result['sections'])
            seo_analysis = analyze_seo_content(full_content, topic)
            logging.info(f"Analyse SEO effectuée: {json.dumps(seo_analysis, indent=2)}")
            
            # Analyse sémantique
            semantic_analysis = analyze_semantic_fields(full_content, wp_db)
            logging.info(f"Analyse sémantique effectuée: {json.dumps(semantic_analysis, indent=2)}")
            
            # Suggestion de liens internes pour chaque section
            for section in result['sections']:
                suggested_links = suggest_internal_links(section['content'], wp_db, topic)
                if suggested_links:
                    # Ajouter les suggestions de liens dans le contenu
                    section['content'] += "\n\nLiens internes suggérés:\n"
                    for link in suggested_links:
                        section['content'] += f"{{{{LIEN_VERS: \"{link['title']}\"}}}}\n"
            
            # Formater le contenu avec une meilleure structure
            formatted_content = f"{result['title']}\n\n"
            formatted_content += f"{result['content']}\n\n"
            
            for section in result['sections']:
                level = section['level']
                title = section['title']
                content = section['content']
                
                # Ajouter le titre avec le bon niveau
                formatted_content += f"{'=' * (level + 1)} {title} {'=' * (level + 1)}\n\n"
                # Ajouter le contenu
                formatted_content += f"{content}\n\n"
            
            result['content'] = formatted_content
            
            # Ajouter les analyses au résultat
            result['seo_analysis'] = seo_analysis
            result['semantic_analysis'] = semantic_analysis
            
            logging.info(f"Page pilier générée avec succès. Titre: {result['title']}")
            return result
            
        except json.JSONDecodeError as e:
            logging.error(f"Erreur de parsing JSON: {str(e)}")
            logging.error(f"Contenu reçu: {content}")
            return None
            
    except Exception as e:
        logging.error(f"Erreur lors de la génération de la page pilier: {str(e)}")
        logging.exception("Trace complète de l'erreur:")
        return None

def analyze_logs(log_content):
    """Analyse le contenu des logs et retourne un résumé des erreurs."""
    try:
        lines = log_content.split('\n')
        error_summary = []
        current_error = None
        
        for line in lines:
            if 'ERROR' in line:
                # Nouvelle erreur détectée
                if current_error:
                    error_summary.append(current_error)
                current_error = {
                    'timestamp': line.split(' - ')[0] if ' - ' in line else 'N/A',
                    'message': line.split('ERROR - ')[-1] if 'ERROR - ' in line else line,
                    'trace': []
                }
            elif current_error and line.strip():
                # Ajouter les lignes de trace à l'erreur courante
                current_error['trace'].append(line.strip())
        
        # Ajouter la dernière erreur si elle existe
        if current_error:
            error_summary.append(current_error)
            
        return error_summary
    except Exception as e:
        return [{'message': f'Erreur lors de l\'analyse des logs: {str(e)}'}]

def display_error_summary(error_summary):
    """Affiche un résumé des erreurs dans l'interface Streamlit."""
    if not error_summary:
        st.info("Aucune erreur n'a été détectée dans les logs.")
        return
        
    st.error("📊 Résumé des erreurs détectées:")
    
    for i, error in enumerate(error_summary, 1):
        with st.expander(f"Erreur #{i} - {error['message']}", expanded=True):
            st.write(f"**Timestamp:** {error['timestamp']}")
            st.write(f"**Message:** {error['message']}")
            if error['trace']:
                st.write("**Trace complète:**")
                st.code('\n'.join(error['trace']))

def save_to_word(pillar_page, topic):
    """Sauvegarde la page pilier dans un fichier Word avec des hyperliens."""
    try:
        # Créer d'abord le document Word de base
        doc = Document()
        
        # Titre principal
        title = doc.add_heading(pillar_page['title'], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenu
        content = pillar_page['content']
        sections = content.split('\n\n')
        
        for section in sections:
            if section.strip():
                # Détecter les niveaux de titre
                if section.startswith('==='):
                    # Titre de niveau 2
                    title_text = section.replace('===', '').strip()
                    doc.add_heading(title_text, level=2)
                elif section.startswith('===='):
                    # Titre de niveau 3
                    title_text = section.replace('====', '').strip()
                    doc.add_heading(title_text, level=3)
                else:
                    # Paragraphe normal
                    p = doc.add_paragraph(section.strip())
                    p.style = 'Normal'
        
        # Sauvegarder le document dans un fichier temporaire
        temp_filename = f"temp_page_pilier_{topic.replace(' ', '_').lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(temp_filename)
        
        try:
            # Ouvrir le fichier temporaire et le traiter pour les hyperliens
            with open(temp_filename, 'rb') as f:
                final_doc = process_word_document(f)
            
            if final_doc:
                # Sauvegarder le fichier final
                filename = f"page_pilier_{topic.replace(' ', '_').lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                # Écrire le contenu du BytesIO dans le fichier
                with open(filename, 'wb') as f:
                    f.write(final_doc.getvalue())
                logging.info(f"Document Word sauvegardé avec succès: {filename}")
                return filename
            else:
                logging.error("Erreur lors du traitement des hyperliens")
                return None
                
        finally:
            # Nettoyer le fichier temporaire
            try:
                os.remove(temp_filename)
            except:
                pass
            
    except Exception as e:
        logging.error(f"Erreur lors de la sauvegarde du fichier Word: {str(e)}")
        logging.exception("Trace complète de l'erreur:")
        return None

def main():
    """Interface principale de l'application Streamlit."""
    st.title("📚 Générateur de Pages Pilier")
    st.write("Générez des pages pilier optimisées basées sur votre contenu existant")
    
    # Informations sur l'application
    with st.expander("ℹ️ Informations"):
        st.write("""
        Cette application:
        - Génère des pages pilier optimisées pour le SEO
        - Utilise uniquement le contenu de votre site existant
        - Crée automatiquement des liens internes pertinents
        - Structure le contenu de manière professionnelle
        - Exporte le résultat en format Word (.docx)
        - Analyse la qualité SEO du contenu généré
        """)
    
    # Champ de saisie pour le sujet
    topic = st.text_input(
        "Sujet de la page pilier",
        help="Entrez le sujet principal de la page pilier que vous souhaitez créer"
    )
    
    if topic:
        # Bouton pour générer la page pilier
        if st.button("🔄 Générer la page pilier", type="primary"):
            with st.spinner("Génération en cours..."):
                try:
                    # Charger la base de données WordPress
                    logging.info("Chargement de la base de données WordPress...")
                    wp_db = load_wordpress_db()
                    if not wp_db:
                        error_msg = "Impossible de charger la base de données WordPress"
                        logging.error(error_msg)
                        st.error(error_msg)
                        return
                    
                    logging.info(f"Base de données WordPress chargée avec {len(wp_db)} articles")
                    
                    # Générer la page pilier
                    pillar_page = generate_pillar_page(topic, wp_db)
                    
                    if pillar_page:
                        # Sauvegarder dans un fichier Word
                        word_file = save_to_word(pillar_page, topic)
                        
                        if word_file:
                            st.success(f"✅ Page pilier générée avec succès! Fichier sauvegardé: {word_file}")
                            
                            # Afficher un aperçu du contenu
                            st.markdown("### Aperçu du contenu")
                            st.markdown(f"**Titre:** {pillar_page['title']}")
                            st.markdown("**Contenu:**")
                            st.text(pillar_page['content'])
                            
                            # Afficher l'analyse SEO
                            st.markdown("### 📊 Analyse SEO")
                            seo_analysis = pillar_page.get('seo_analysis', {})
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.metric("Densité du mot-clé", f"{seo_analysis.get('keyword_density', 0)}%")
                                st.metric("Score de lisibilité", f"{seo_analysis.get('readability_score', 0)}/100")
                                st.metric("Nombre total de mots", seo_analysis.get('total_words', 0))
                            
                            with col2:
                                st.metric("Phrases trop longues", seo_analysis.get('long_sentences', 0))
                                if seo_analysis.get('vague_expressions'):
                                    st.warning("⚠️ Expressions vagues détectées:")
                                    for expr in seo_analysis['vague_expressions']:
                                        st.write(f"- {expr}")
                            
                            # Afficher l'analyse sémantique
                            st.markdown("### 🔍 Analyse Sémantique")
                            semantic_analysis = pillar_page.get('semantic_analysis', {})
                            
                            if semantic_analysis.get('top_keywords'):
                                st.write("**Mots-clés principaux:**")
                                keywords_df = pd.DataFrame(semantic_analysis['top_keywords'], 
                                                         columns=['Mot-clé', 'Fréquence'])
                                st.dataframe(keywords_df)
                            
                            if semantic_analysis.get('semantic_groups'):
                                st.write("**Groupes sémantiques:**")
                                for word, related in semantic_analysis['semantic_groups'].items():
                                    with st.expander(f"Groupe: {word}"):
                                        for related_word, similarity in related:
                                            st.write(f"- {related_word} (similarité: {similarity:.2f})")
                            
                            # Afficher les liens internes suggérés
                            st.markdown("### 🔗 Liens Internes Suggérés")
                            for section in pillar_page.get('sections', []):
                                if "Liens internes suggérés:" in section['content']:
                                    st.write(f"**Section: {section['title']}**")
                                    links = re.findall(r'{{LIEN_VERS: "([^"]+)"}}', section['content'])
                                    for link in links:
                                        st.write(f"- {link}")
                        else:
                            st.error("❌ Erreur lors de la sauvegarde du fichier Word")
                    else:
                        error_msg = "❌ Erreur lors de la génération de la page pilier. Consultez les logs pour plus de détails."
                        logging.error(error_msg)
                        st.error(error_msg)
                        
                        # Afficher l'analyse des logs
                        log_content = get_latest_log_content()
                        error_summary = analyze_logs(log_content)
                        display_error_summary(error_summary)
                        
                        # Afficher les logs complets
                        with st.expander("📋 Logs complets"):
                            st.code(log_content)
                    
                except Exception as e:
                    error_msg = f"❌ Une erreur s'est produite: {str(e)}"
                    logging.error(error_msg)
                    logging.exception("Trace complète de l'erreur:")
                    st.error(error_msg)
                    
                    # Afficher l'analyse des logs
                    log_content = get_latest_log_content()
                    error_summary = analyze_logs(log_content)
                    display_error_summary(error_summary)
                    
                    # Afficher les logs complets
                    with st.expander("📋 Logs complets"):
                        st.code(log_content)

if __name__ == "__main__":
    main()